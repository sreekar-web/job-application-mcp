from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from ats.keyword_extractor import extract_keywords

def build_resume_docx(
    output_path,
    role_variant: dict,
    job_description: str = None,
    highlight_keywords: bool = True
):
    """
    Build an ATS-safe DOCX resume from a role variant definition.
    
    Args:
        output_path: Path to save the DOCX file
        role_variant: Role variant dict with resume_sections
        job_description: Optional JD to extract keywords for highlighting
        highlight_keywords: Whether to highlight matched keywords
    
    Returns:
        dict with success status, output_path, and any warnings
    """
    
    try:
        # Extract resume sections from variant
        sections = role_variant.get("resume_sections", {})
        if not sections:
            return {
                "success": False,
                "error": f"No resume_sections found in role variant"
            }
        
        # Create document
        doc = Document()
        
        # Set default font for ATS compatibility
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # --- SUMMARY ---
        if "summary" in sections:
            doc.add_heading("SUMMARY", level=1)
            doc.add_paragraph(sections["summary"])
        
        # --- SKILLS ---
        if "allowed_skills" in role_variant:
            doc.add_heading("SKILLS", level=1)
            skills = role_variant.get("allowed_skills", [])
            doc.add_paragraph(", ".join(skills))
        
        # --- EXPERIENCE ---
        if "experience" in sections:
            doc.add_heading("EXPERIENCE", level=1)
            for job in sections["experience"]:
                title = job.get("title", "")
                company = job.get("company", "")
                duration = job.get("duration", "")
                
                doc.add_paragraph(
                    f"{title} – {company} ({duration})",
                    style='Heading 2'
                )
                
                for resp in job.get("responsibilities", []):
                    doc.add_paragraph(f"• {resp}", style='List Bullet')
        
        # --- EDUCATION ---
        if "education" in sections:
            doc.add_heading("EDUCATION", level=1)
            for edu in sections["education"]:
                doc.add_paragraph(edu)
        
        # --- CERTIFICATIONS (if present) ---
        if "certifications" in sections:
            certs = sections["certifications"]
            if certs:
                doc.add_heading("CERTIFICATIONS", level=1)
                for cert in certs:
                    doc.add_paragraph(f"• {cert}", style='List Bullet')
        
        # --- KEY PROJECTS (if present) ---
        if "key_projects" in sections:
            projects = sections["key_projects"]
            if projects:
                doc.add_heading("KEY PROJECTS", level=1)
                for project in projects:
                    doc.add_paragraph(f"• {project}", style='List Bullet')
        
        # Save document
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        
        # Extract and log matched keywords if JD provided
        matched_keywords = []
        if job_description and highlight_keywords:
            allowed_skills = role_variant.get("allowed_skills", [])
            matched_keywords = extract_keywords(job_description, allowed_skills)
        
        return {
            "success": True,
            "output_path": str(output_path),
            "matched_keywords": matched_keywords,
            "warnings": []
        }
    
    except KeyError as e:
        return {
            "success": False,
            "error": f"Missing required field in role variant: {e}",
            "warnings": ["Ensure role variant has 'resume_sections' with required content"]
        }
    
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "warnings": ["Check role variant structure and resume_sections format"]
        }

