import re
from pathlib import Path
from docx import Document

# ATS-unsafe elements
FORBIDDEN_TAGS = ["table", "<w:tbl", "column", "header", "footer", "image", "diagram"]
FORBIDDEN_CHARACTERS = ["©", "®", "™", "†", "‡"]

REQUIRED_SECTIONS = [
    "SUMMARY",
    "SKILLS",
    "EXPERIENCE",
    "EDUCATION"
]

def validate_resume_text(text: str) -> dict:
    """Validate resume text content for ATS safety."""
    issues = []
    
    # Check required sections
    upper = text.upper()
    for section in REQUIRED_SECTIONS:
        if section not in upper:
            issues.append(f"Missing required section: {section}")
    
    # Check forbidden tags (XML/HTML that may confuse ATS)
    for bad_tag in FORBIDDEN_TAGS:
        if bad_tag.lower() in text.lower():
            issues.append(f"ATS violation: Contains '{bad_tag}' (tables/images not ATS-safe)")
    
    # Check forbidden characters (unicode symbols)
    for char in FORBIDDEN_CHARACTERS:
        if char in text:
            issues.append(f"ATS violation: Contains special character '{char}' (not machine-readable)")
    
    # Check for excessive special characters (may indicate encoding issues)
    special_char_count = len(re.findall(r'[^\w\s\-.,;:\'\"]', text))
    if special_char_count > 20:
        issues.append(f"Too many special characters ({special_char_count}) — may confuse ATS parser")
    
    # Check content length
    if len(text.strip()) < 200:
        issues.append("Resume too short — may lack required content for ATS parsing")
    
    if len(text) > 50000:
        issues.append("Resume too long — may cause parsing issues")
    
    # Check for multiple spaces/tabs (indicate formatting cruft)
    if re.search(r'  {2,}', text):
        issues.append("Multiple consecutive spaces detected — clean up formatting")
    
    return {
        "valid": len(issues) == 0,
        "issues": issues
    }

def validate_resume_docx(docx_path: str) -> dict:
    """Validate DOCX file for ATS compatibility."""
    issues = []
    
    try:
        doc = Document(docx_path)
        
        # Extract all text
        text_content = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        
        # Check for tables (major ATS issue)
        if len(doc.tables) > 0:
            issues.append(f"Found {len(doc.tables)} table(s) — tables are not ATS-safe, use bullet points instead")
        
        # Check for images
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref.lower():
                issues.append("Resume contains image(s) — images are not ATS-parseable")
        
        # Check font consistency
        fonts = set()
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if run.font.name:
                    fonts.add(run.font.name)
        
        # Common ATS-safe fonts
        safe_fonts = {"Calibri", "Arial", "Times New Roman", "Courier New", None}
        unsafe_fonts = fonts - safe_fonts
        if unsafe_fonts:
            issues.append(f"Unsafe fonts detected: {', '.join(unsafe_fonts)} — use Calibri, Arial, or Times New Roman")
        
        # Check text content
        text_validation = validate_resume_text(text_content)
        issues.extend(text_validation["issues"])
        
        return {
            "valid": len(issues) == 0,
            "issues": issues,
            "font_count": len(fonts),
            "table_count": len(doc.tables),
            "character_count": len(text_content)
        }
    
    except Exception as e:
        return {
            "valid": False,
            "issues": [f"Failed to validate DOCX: {str(e)}"],
            "error": str(e)
        }

def validate_resume(text_or_path):
    """
    Validate resume (text or DOCX path).
    
    Args:
        text_or_path: Resume text string or path to .docx file
    
    Returns:
        dict with 'valid' (bool) and 'issues' (list of strings)
    """
    if isinstance(text_or_path, str) and text_or_path.endswith(".docx"):
        return validate_resume_docx(text_or_path)
    else:
        return validate_resume_text(str(text_or_path))

